from pathlib import Path
import re
import shutil
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(__file__).resolve().parent
ROOT = BASE.parent
PRIMARY = RGBColor(0x01, 0x42, 0x6A)
SECONDARY = RGBColor(0x00, 0x67, 0xA0)
TEAL = RGBColor(0x00, 0x7B, 0x5F)
GOLD = RGBColor(0xFF, 0xCD, 0x00)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def clean_inline(text):
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text


def add_rich_paragraph(doc, text, style=None):
    para = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = para.add_run(part.strip("`").strip("*"))
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(10)
    return para


def add_table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r"^\s*\|?\s*:?-{3,}", line):
            continue
        cells = [clean_inline(c.strip()) for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = text
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = PRIMARY
    doc.add_paragraph()


def add_image(doc, line):
    match = re.search(r"!\[(.*?)\]\((.*?)\)", line)
    if not match:
        return False
    alt, rel = match.groups()
    img = (BASE / rel).resolve()
    if not img.exists():
        img = (ROOT / rel).resolve()
    if img.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(img), width=Inches(6.2))
        caption = doc.add_paragraph(alt)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(9)
        return True
    add_rich_paragraph(doc, f"[Imagem nao localizada: {alt}]")
    return True


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, PRIMARY, 16, 8),
        ("Heading 2", 13, SECONDARY, 12, 6),
        ("Heading 3", 12, TEAL, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc, title):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "TRT10 - Artefatos de Contratacao MPLS + SD-WAN"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = PRIMARY
    footer = section.footer.paragraphs[0]
    footer.text = title
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build(md_path):
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    setup_styles(doc)
    title = md_path.stem.replace("_", " ")
    add_header_footer(doc, title)

    table_buf = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if table_buf and not line.strip().startswith("|"):
            add_table(doc, table_buf)
            table_buf = []
        if not line.strip():
            continue
        if line.strip().startswith("|"):
            table_buf.append(line)
            continue
        if add_image(doc, line):
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(clean_inline(line[2:]))
            p.style = "Title"
            p.runs[0].font.color.rgb = PRIMARY
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.bold = True
        elif line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
        elif re.match(r"^\d+\.\s+", line):
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s+", "", line), style="List Number")
        elif line.startswith("- "):
            add_rich_paragraph(doc, line[2:], style="List Bullet")
        else:
            add_rich_paragraph(doc, line)
    if table_buf:
        add_table(doc, table_buf)

    out = md_path.with_suffix(".docx")
    try:
        doc.save(out)
    except PermissionError:
        out = md_path.with_name(md_path.stem + "_atualizado.docx")
        doc.save(out)
    return out


def convert_pdf(docx_path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    pdf = docx_path.with_suffix(".pdf")
    return pdf if pdf.exists() else None


if __name__ == "__main__":
    files = [
        BASE / "DFD_MPLS_TRT10.md",
        BASE / "ETP_MPLS_TRT10.md",
        BASE / "TR_MPLS_TRT10.md",
        BASE / "ANEXOS_TR_MPLS_TRT10.md",
        BASE / "PENDENCIAS_E_AJUSTES_MPLS.md",
    ]
    for md in files:
        docx = build(md)
        pdf = convert_pdf(docx)
        print(f"generated {docx}" + (f" and {pdf}" if pdf else ""))
